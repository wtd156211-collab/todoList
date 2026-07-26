from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("output/doc/Flowlist_微信小程序_MVP_开发计划.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "1F2937"
BLUE = "2563EB"
PALE_BLUE = "EFF6FF"
SLATE = "475569"
OFF_WHITE = "F8FAFC"
LINE = "CBD5E1"
ORANGE = "C2410C"


def font(run, size=10, color=NAVY, bold=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def shade(cell, value):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), value)


def border(cell, value=LINE):
    props = cell._tc.get_or_add_tcPr()
    root = props.first_child_found_in("w:tcBorders")
    if root is None:
        root = OxmlElement("w:tcBorders")
        props.append(root)
    for edge in ("top", "left", "bottom", "right"):
        child = root.find(qn(f"w:{edge}"))
        if child is None:
            child = OxmlElement(f"w:{edge}")
            root.append(child)
        child.set(qn("w:val"), "single")
        child.set(qn("w:sz"), "4")
        child.set(qn("w:space"), "0")
        child.set(qn("w:color"), value)


def cell_text(cell, value, size=8.8, bold=False, color=NAVY):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(str(value))
    font(r, size, color, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    border(cell)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    row = t.rows[0]
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)
    for i, title in enumerate(headers):
        row.cells[i].width = Cm(widths[i])
        shade(row.cells[i], NAVY)
        cell_text(row.cells[i], title, 8.8, True, "FFFFFF")
    for n, values in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(values):
            cells[i].width = Cm(widths[i])
            shade(cells[i], "FFFFFF" if n % 2 == 0 else OFF_WHITE)
            cell_text(cells[i], value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def heading(doc, value, level=1):
    p = doc.add_heading(value, level=level)
    p.paragraph_format.space_before = Pt(15 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p


def body(doc, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.23
    r = p.add_run(value)
    font(r, 10.1)
    return p


def bullet(doc, value, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(value)
    font(r, 9.8)
    return p


def note(doc, title, value, color=BLUE):
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    shade(c, PALE_BLUE)
    border(c, color)
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + "  ")
    font(r, 10, color, True)
    r = p.add_run(value)
    font(r, 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Flowlist | 微信小程序 MVP 开发计划 | ")
    font(r, 8, SLATE)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.7)
sec.bottom_margin = Cm(1.7)
sec.left_margin = Cm(1.85)
sec.right_margin = Cm(1.85)
footer(sec)

styles = doc.styles
for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
    st = styles[name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.1)
styles["Title"].font.size = Pt(29)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor.from_string(NAVY)
for name, size, color in [("Heading 1", 17, NAVY), ("Heading 2", 12.5, BLUE), ("Heading 3", 11, SLATE)]:
    styles[name].font.size = Pt(size)
    styles[name].font.bold = True
    styles[name].font.color.rgb = RGBColor.from_string(color)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(78)
r = p.add_run("FLOWLIST")
font(r, 13, BLUE, True)
p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(16)
p.add_run("微信小程序 MVP\n开发计划")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("个人任务管理 · 原生微信小程序 + 腾讯云 CloudBase\n依据 Flowlist UI / UX Wireframe Spec 编制")
font(r, 11.5, SLATE)
doc.add_paragraph()

t = doc.add_table(rows=5, cols=2)
for row, pair in zip(t.rows, [
    ("文档版本", "v1.0"),
    ("编制日期", "2026-07-26"),
    ("目标平台", "微信小程序"),
    ("范围", "个人任务管理 MVP，不含协作能力"),
    ("交付目标", "可提交审核、可稳定使用、可发送订阅提醒的任务管理小程序"),
]):
    row.cells[0].width = Cm(4.0)
    row.cells[1].width = Cm(11.0)
    shade(row.cells[0], NAVY)
    cell_text(row.cells[0], pair[0], 10, True, "FFFFFF")
    shade(row.cells[1], "FFFFFF")
    cell_text(row.cells[1], pair[1], 10)
doc.add_page_break()

# 1
heading(doc, "1. 项目目标、范围与关键决策")
body(doc, "Flowlist 首版定位为微信内即可使用的个人任务管理工具。用户可快速建立任务、设置时间和优先级、在日历中查看安排，并在获得授权后通过微信服务通知接收到期提醒。技术方案以微信生态为唯一目标，避免为 App 或 Web 的跨端兼容增加首版复杂度。")
table(doc, ["目标", "首版完成定义"], [
    ("低门槛进入", "打开小程序后以微信身份初始化账户，不出现邮箱注册或密码登录流程。"),
    ("快速完成任务闭环", "支持创建、编辑、完成、删除、搜索、分类与日历查看。"),
    ("可控提醒", "用户为具体提醒主动订阅后，由定时云函数发送到期服务通知。"),
    ("数据私有", "所有任务、附件、通知仅由其 OpenID 所属用户访问。"),
    ("可审核发布", "具备隐私说明、错误处理、真机验证和微信公众平台审核材料。"),
], [4.0, 11.0])
heading(doc, "1.1 纳入 MVP", 2)
for text in [
    "微信身份初始化、头像和昵称的可选完善。",
    "首页：今日 / 即将到来 / 全部、关键词搜索、任务完成、快速新增。",
    "任务：标题、备注、截止时间、优先级、分类、提醒、子任务、图片或 PDF 附件。",
    "日历：按月查看、任务日期标识、按日任务列表。",
    "通知中心：到期提醒、完成记录、已读状态和清除全部。",
    "我的：资料、分类管理、提醒说明和隐私入口。",
]:
    bullet(doc, text)
heading(doc, "1.2 不纳入 MVP", 2)
for text in [
    "邮箱、Google、Apple 登录，以及独立账号密码体系。",
    "任务共享、成员邀请、评论、@提及和协作权限。",
    "循环任务、模板、统计报表、Web 端和原生 App。",
    "深色模式、多语言、离线冲突合并与回收站。",
]:
    bullet(doc, text)
note(doc, "首要前置条件", "开工第 0 周必须在微信公众平台确认小程序主体、服务类目和“任务到期提醒”可用的订阅消息模板；此项决定外部提醒能否纳入发布承诺。", ORANGE)

# 2
heading(doc, "2. 线框图到小程序的页面改造")
table(doc, ["原页面", "小程序实现", "差异与验收"], [
    ("欢迎 / 启动", "保留品牌、价值主张和“开始使用”。", "读取登录态；首次使用进入微信身份初始化，老用户直达首页。"),
    ("注册 / 登录", "合并为微信快速开始页。", "移除邮箱、密码、Google、Apple；不强制读取头像昵称。"),
    ("通知授权", "改为提醒说明与任务保存时的订阅授权。", "不能作为一次性系统权限；用户拒绝后任务仍可创建。"),
    ("首页 / 任务列表", "保持原布局和底部导航。", "支持搜索、防抖、筛选、空状态和一键完成。"),
    ("任务详情 / 创建", "保持字段；使用原生日期时间选择与自定义优先级、分类组件。", "标题必填；删除二次确认；提醒有订阅状态提示。"),
    ("日历 / 通知 / 我的", "保留三页；通知页以站内记录为唯一可信来源。", "点击通知跳任务详情；分类修改即时反映到创建任务页。"),
], [3.0, 5.5, 6.5])
heading(doc, "2.1 默认业务规则", 2)
table(doc, ["规则", "默认值"], [
    ("任务优先级", "低 / 中 / 高，创建默认“中”。"),
    ("分类", "内置 Work、Personal；可新建、改名、改色。被任务引用的分类不可直接删除。"),
    ("提醒选项", "准时、提前 10 分钟、30 分钟、1 小时、1 天；默认关闭。"),
    ("附件", "图片与 PDF；单文件上限 10 MB、每任务最多 5 个。"),
    ("时区", "服务端存 UTC，任务记录 IANA 时区，界面按任务设定时区展示。"),
    ("删除", "确认后删除任务、子任务、通知和附件元数据；云端文件异步清理。"),
], [3.8, 11.2])

# 3
heading(doc, "3. 总体技术架构")
body(doc, "首版使用原生小程序，不采用跨端框架。原因是产品只运行于微信，原生 API 对登录、订阅消息、文件选择、云开发和审核环境的兼容性最佳；若后续确认需要 H5 或 App，再单独评估 Taro 的迁移或重构。")
table(doc, ["层级", "选型", "职责"], [
    ("页面层", "原生微信小程序 + TypeScript + WXML + WXSS / Less", "实现九个主页面、路由、交互、空/错/加载状态。"),
    ("设计组件", "自定义基础组件；Vant Weapp 仅作 Dialog、Popup、Picker 等基础能力", "保持线框稿的定制视觉，避免组件库主导页面样式。"),
    ("状态与请求", "MobX-miniprogram + Promise API 封装", "保存筛选、页面状态和请求生命周期；业务数据以云端为准。"),
    ("业务服务", "CloudBase 云函数，Node.js + TypeScript", "统一身份归属、任务 CRUD、附件校验、提醒调度与日志。"),
    ("数据", "CloudBase 文档数据库", "存储用户、分类、任务和站内通知。"),
    ("文件", "CloudBase 云存储", "存储任务图片/PDF；文件路径由云函数创建并校验。"),
    ("提醒", "订阅消息 + 定时云函数", "扫描应提醒任务，创建站内通知，发送一次订阅消息。"),
    ("质量与发布", "微信开发者工具、自动化脚本、体验版、日志告警", "覆盖开发、真机验证、审核和线上问题定位。"),
], [3.0, 5.3, 6.7])
heading(doc, "3.1 运行流", 2)
table(doc, ["触发", "处理", "结果"], [
    ("用户打开小程序", "调用初始化云函数；云端由上下文识别 OpenID 并创建/读取 users 记录。", "恢复用户资料、默认分类和首页数据。"),
    ("用户保存任务", "前端校验字段；若开启提醒，先请求订阅；再调用 saveTask。", "任务、子任务、附件元数据和提醒状态原子保存。"),
    ("定时任务运行", "每 5 分钟扫描到期窗口，通过原子状态锁定待发提醒。", "写 notifications 后下发订阅消息；失败留日志供重试。"),
    ("点击服务通知", "携带 taskId 打开小程序，路由跳转任务详情。", "通知标记已读；任务不存在时进入通知中心。"),
], [3.2, 8.4, 3.4])

# 4
heading(doc, "4. 代码目录、数据模型与云函数")
heading(doc, "4.1 推荐目录", 2)
table(doc, ["路径", "内容"], [
    ("miniprogram/pages", "welcome、home、task-form、task-detail、calendar、notifications、profile 与设置子页。"),
    ("miniprogram/components", "task-card、priority-selector、category-chip、empty-state、confirm-dialog、bottom-nav。"),
    ("miniprogram/services", "cloud.ts、auth.ts、task.ts、notification.ts、upload.ts；封装云函数调用。"),
    ("miniprogram/store", "session、task-filter、draft；不在本地长期复制业务数据。"),
    ("miniprogram/utils", "日期/时区、表单校验、错误映射、节流与日志工具。"),
    ("cloudfunctions", "bootstrapUser、taskList、taskSave、taskToggle、taskDelete、fileSign、notifications、sendDueReminders。"),
    ("shared", "任务 DTO、枚举、Zod 校验规则、通知 payload 类型。"),
], [4.3, 10.7])
heading(doc, "4.2 集合与字段", 2)
table(doc, ["集合", "关键字段", "设计说明"], [
    ("users", "_id, openid, nickname, avatarFileId, timezone, createdAt", "openid 只由云函数读取和写入，不回传至页面。"),
    ("categories", "_id, ownerOpenid, name, color, isSystem, sortOrder", "初始化 Work / Personal；同一用户名称不可重复。"),
    ("tasks", "_id, ownerOpenid, title, note, status, priority, categoryId, dueAtUtc, timezone, reminder, subtasks, attachments", "子任务与附件元数据嵌入，首版避免多集合联查。"),
    ("notifications", "_id, ownerOpenid, taskId, type, title, content, isRead, createdAt", "通知中心数据源；和订阅消息投递记录关联。"),
    ("reminderDeliveries", "taskId, triggerAtUtc, status, attemptCount, sentAt, error", "以 taskId + triggerAtUtc 建唯一约束，实现幂等发送与失败追踪。"),
], [3.1, 7.6, 4.3])
heading(doc, "4.3 云函数接口契约", 2)
table(doc, ["函数", "输入", "输出 / 校验"], [
    ("bootstrapUser", "可选昵称、头像引用", "返回用户和默认分类；只以云端 OpenID 作为归属。"),
    ("taskList", "scope, keyword, cursor, selectedDate", "返回分页任务；仅查询当前用户的记录。"),
    ("taskSave", "TaskDraft", "校验标题、时间、附件限制、分类归属；返回完整任务。"),
    ("taskToggle / taskDelete", "taskId, action", "验证归属；更新完成状态或级联清理关联记录。"),
    ("fileSign / fileCleanup", "任务草稿或 taskId、文件元数据", "限制 MIME/大小/数量，生成私有文件路径并清理孤儿文件。"),
    ("notificationList / notificationRead", "cursor 或 notificationId", "仅操作当前用户通知。"),
    ("sendDueReminders", "定时器 event", "仅允许定时器调用；原子领取提醒后发送消息并记录回执。"),
], [3.1, 5.0, 6.9])

# 5
heading(doc, "5. 微信身份、权限与安全方案")
body(doc, "小程序不保留传统注册表单。页面首次进入时使用微信登录码触发云端身份初始化；服务端依平台流程识别用户身份。昵称和头像均为可选资料，不能成为使用任务功能的前置条件。")
table(doc, ["场景", "实现规则", "验收点"], [
    ("身份识别", "云函数从可信上下文取得 OpenID；客户端不传 ownerOpenid。", "修改请求参数不能读取或修改其他用户任务。"),
    ("数据库访问", "客户端不直接执行敏感集合 CRUD；业务读写全部经云函数。", "云函数逐次以当前 OpenID 加入查询条件。"),
    ("文件访问", "附件存入用户/任务隔离路径；下载地址按需、短时生成。", "猜测 fileId 或 taskId 不能获得他人附件。"),
    ("输入安全", "服务端复用校验规则，限制字段长度、MIME、文件大小与数量。", "前端绕过校验仍会被云端拒绝。"),
    ("隐私", "最小化收集，不请求无关手机号、通讯录、位置或用户资料。", "隐私政策、收集清单、删除说明与实际 API 一致。"),
], [3.0, 8.0, 4.0])
note(doc, "安全红线", "不要在小程序代码、云函数日志或接口响应中暴露 AppSecret、session_key 或用户 OpenID 列表。")

# 6
heading(doc, "6. 订阅消息与站内提醒设计")
body(doc, "订阅消息是小程序首版唯一的外部提醒路径。用户不能在启动页一次性永久开启所有提醒；应在他主动保存“带提醒的任务”时解释用途并请求该模板授权。拒绝授权不阻止任务保存，但必须在详情和设置页展示“未开启微信提醒”。")
table(doc, ["步骤", "前端行为", "云端行为"], [
    ("1. 开启提醒", "用户选择提醒时间并点击创建/保存。", "暂不创建可发送提醒，等待授权结果。"),
    ("2. 请求订阅", "调用订阅消息授权；展示模板所表达的到期提醒用途。", "记录该任务的授权结果与模板 ID。"),
    ("3. 保存任务", "无论授权同意或拒绝，均保存任务；拒绝时给出启用说明。", "将 reminder 标为 pending 或 skipped。"),
    ("4. 到期扫描", "无前端参与。", "每 5 分钟扫描 triggerAtUtc 窗口，以唯一投递记录领取任务。"),
    ("5. 消息点击", "通过启动参数读取 taskId，跳转详情。", "标记通知已读，记录跳转日志。"),
], [2.5, 7.0, 5.5])
heading(doc, "6.1 可靠性规则", 2)
for text in [
    "提醒调度按 5 分钟窗口运行，产品文案不承诺秒级准确；任务显示的时间仍以用户设定时刻为准。",
    "发送前写入 reminderDeliveries 的唯一记录；重复执行定时器时只允许一个实例投递。",
    "消息接口失败时记录错误和尝试次数；无效授权、模板变更等情况标记为不可重试并在站内通知页可见。",
    "用户未授权、授权失效或模板不可用时，降级为首页/通知中心的站内提醒，不宣称已发送微信提醒。",
]:
    bullet(doc, text)
note(doc, "验证门槛", "订阅消息模板、授权弹窗、发送结果和点击跳转必须使用真实微信账号与体验版测试；开发者工具模拟器不能作为最终验收依据。", ORANGE)

# 7
heading(doc, "7. 开发迭代与交付物")
body(doc, "以下按一名全栈开发者、每周一个迭代估算，共 5 周。设计确认、账号资质和订阅模板审核不应压缩到最后一周。")
table(doc, ["迭代", "实施内容", "交付与验收"], [
    ("第 0 周\n准备", "确认类目/模板、创建 AppID 和 CloudBase 环境、整理设计 token、定义数据与隐私清单。", "技术基线、环境变量、页面清单、模板可用性结论。"),
    ("第 1 周\n基础", "小程序工程、设计组件、云函数骨架、身份初始化、分类、首页导航。", "体验版可登录，首页/我的可用，数据按用户隔离。"),
    ("第 2 周\n任务闭环", "任务创建、详情编辑、完成、删除、搜索、筛选、校验与错误状态。", "个人任务 CRUD 完整通过，首页刷新一致。"),
    ("第 3 周\n日历与文件", "月历、按日列表、子任务、附件、设置、站内通知中心。", "全部主页面可走通，文件和分类限制有效。"),
    ("第 4 周\n提醒", "订阅授权、定时器、投递幂等、通知深链、失败降级。", "双真机账号完成授权、发送、点击打开任务的全链路验证。"),
    ("第 5 周\n发布", "回归、兼容性、性能、隐私、审核材料、体验版灰度和线上监控。", "Release Candidate、测试报告、审核包和回滚说明。"),
], [2.6, 8.2, 4.2])
heading(doc, "7.1 每个迭代的完成标准", 2)
for text in [
    "功能在真机与开发者工具均可运行，包含加载、空状态、网络失败和权限拒绝状态。",
    "关键云函数具备输入校验、身份归属校验、结构化日志和可读错误码。",
    "每周演示时逐条对照线框图与本计划的验收标准；新增需求进入 Backlog，不插入当前迭代。",
]:
    bullet(doc, text)

# 8
heading(doc, "8. 测试、审核与发布清单")
table(doc, ["测试层", "覆盖内容", "通过条件"], [
    ("单元", "日期/时区、提醒窗口、表单校验、筛选和优先级排序。", "关键规则有自动化断言，修改后可重复执行。"),
    ("云函数", "OpenID 归属、越权访问、附件限制、删除级联、提醒幂等。", "伪造 taskId/openid 无法访问；同提醒最多投递一次。"),
    ("端到端", "首次使用、创建任务、完成、日历、设置、退出/重进。", "关键路径无阻断，错误后可恢复。"),
    ("真实设备", "安卓/iOS 微信、订阅授权、服务通知、点击深链、文件选择。", "至少各一台真实设备完成记录和截图。"),
    ("审核前", "隐私政策、类目、账号主体、服务类目、订阅模板、反馈入口。", "公众平台配置与实际收集/使用数据一致。"),
], [2.8, 7.4, 4.8])
heading(doc, "8.1 发布验收清单", 2)
for text in [
    "不同微信账号之间完全隔离任务、通知和附件；所有越权请求被拒绝。",
    "用户拒绝订阅提醒时仍可创建任务，且不会出现“已开启提醒”的误导状态。",
    "已授权的测试任务在预定窗口内仅收到一次服务通知；点击后打开正确任务。",
    "离线或服务异常时，不丢失已成功保存的任务；失败操作有可理解的重试提示。",
    "首页、日历、任务详情、通知和我的页面在常见手机屏幕上无截断、无遮挡。",
    "隐私说明、用户协议、客服/反馈入口、图标、名称和审核测试账号已准备。",
]:
    bullet(doc, text)

# 9
heading(doc, "9. 风险、降级策略与后续 Backlog")
table(doc, ["风险", "影响", "策略"], [
    ("订阅模板不可用或审核受限", "外部到期提醒不可承诺。", "第 0 周确认；降级为站内通知与首页提醒，并从宣传文案移除“微信提醒”。"),
    ("用户拒绝订阅", "部分任务不会发送服务通知。", "任务照常保存；详情页显示未开启状态，允许用户下次保存提醒时再次授权。"),
    ("定时任务延迟或失败", "提醒可能晚到或漏发。", "5 分钟扫描、幂等投递、错误日志、有限重试和站内通知兜底。"),
    ("附件成本增长", "云存储成本与审核风险增加。", "限制 MIME、单文件 10 MB、每任务 5 个，定期清理孤儿文件。"),
    ("范围膨胀", "MVP 延期。", "协作、循环任务、深色模式、统计全部入 Backlog，单独评估。"),
], [3.6, 5.2, 6.2])
heading(doc, "9.1 后续版本候选", 2)
table(doc, ["优先级", "能力", "启动条件"], [
    ("P1", "循环任务、任务模板、归档/回收站", "首版留存和日常任务使用稳定后。"),
    ("P1", "深色模式、多语言、桌面/Web 端", "明确用户设备分布和目标市场后。"),
    ("P2", "共享任务、成员、评论和 @提及", "独立完成协作角色、邀请、权限和通知设计后。"),
    ("P2", "统计、习惯追踪、AI 任务拆解", "有足够的真实使用数据并明确隐私边界后。"),
], [2.3, 6.8, 5.9])
note(doc, "结论", "本方案将原线框图完整收敛为可审核的个人任务管理微信小程序。核心风险集中于订阅消息模板和真实设备提醒链路，应在第 0 周和第 4 周分别完成准入确认与全链路验收。")

doc.save(OUT)
print(OUT.resolve())
