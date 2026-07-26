from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("output/doc/Flowlist_个人任务管理_MVP_开发计划.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "1F2937"
BLUE = "2563EB"
LIGHT_BLUE = "EFF6FF"
SLATE = "475569"
LIGHT_GRAY = "F8FAFC"
BORDER = "CBD5E1"
RED = "B91C1C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, bold=False, color=NAVY, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        if widths:
            hdr.cells[i].width = Cm(widths[i])
        set_cell_shading(hdr.cells[i], NAVY)
        set_cell_text(hdr.cells[i], header, bold=True, color="FFFFFF", size=9)
    for index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if widths:
                cells[i].width = Cm(widths[i])
            set_cell_shading(cells[i], "FFFFFF" if index % 2 == 0 else LIGHT_GRAY)
            set_cell_text(cells[i], value, size=8.7)
    doc.add_paragraph()
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.55 + level * 0.45)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10)
    return p


def add_note(doc, title, text, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(cell, color)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(7)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.3)
    return p


def add_page_break(doc):
    doc.add_page_break()


def set_run_font(run, size=None, color=None, bold=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Flowlist  |  个人任务管理 MVP 开发计划  |  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.7)
section.left_margin = Cm(1.9)
section.right_margin = Cm(1.9)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.3)
styles["Title"].font.name = "Microsoft YaHei"
styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Title"].font.size = Pt(30)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor.from_string(NAVY)
for style_name, size, color in [("Heading 1", 17, NAVY), ("Heading 2", 12.5, BLUE), ("Heading 3", 11, SLATE)]:
    style = styles[style_name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)

for s in doc.sections:
    footer = s.footer.paragraphs[0]
    set_page_number(footer)
    for run in footer.runs:
        set_run_font(run, 8, SLATE)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(88)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FLOWLIST")
set_run_font(r, 13, BLUE, True)
p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(9)
p.paragraph_format.space_after = Pt(16)
p.add_run("个人任务管理 MVP\n开发计划")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("基于《Flowlist UI / UX Wireframe Spec》\n范围版本：个人任务管理首版")
set_run_font(r, 12, SLATE)

doc.add_paragraph()
table = doc.add_table(rows=4, cols=2)
table.autofit = False
cover_rows = [("文档版本", "v1.0"), ("编制日期", "2026-07-26"), ("目标平台", "iOS / Android"), ("交付目标", "可发布、可提醒、可维护的个人任务管理应用")]
for row, (k, v) in zip(table.rows, cover_rows):
    row.cells[0].width = Cm(4.2)
    row.cells[1].width = Cm(10.5)
    set_cell_shading(row.cells[0], NAVY)
    set_cell_text(row.cells[0], k, bold=True, color="FFFFFF", size=10)
    set_cell_shading(row.cells[1], "FFFFFF")
    set_cell_text(row.cells[1], v, size=10)

add_page_break(doc)

# 1 Overview
add_heading(doc, "1. 项目目标与交付边界")
add_body(doc, "Flowlist 是一款移动端个人任务管理应用。首版的重点不是堆叠复杂功能，而是让用户可以快速记录任务、在合适的时间获得提醒，并在首页和日历中清楚地掌握当天计划。")
add_table(doc, ["目标", "首版交付定义"], [
    ("快速记录", "用户可在 1 分钟内创建含标题、日期、优先级、分类和提醒的任务。"),
    ("每日执行", "首页展示今日、即将到来和全部任务；支持搜索、筛选与一键完成。"),
    ("可靠提醒", "任务到期前按用户设定的时间发送本地/远程提醒，并可点击直达任务。"),
    ("账户与数据安全", "支持邮箱、Google、Apple 登录；任务数据按用户隔离。"),
    ("发布准备", "具备错误处理、空状态、真机推送测试、隐私说明和商店构建能力。")
], [4.1, 10.6])
add_heading(doc, "1.1 明确纳入 MVP 的范围", 2)
for text in [
    "欢迎引导、注册、登录和登录态恢复。",
    "任务创建、查看、编辑、完成、删除，及子任务、备注、附件。",
    "首页任务列表、关键词搜索、今日/即将到来/全部筛选。",
    "按月查看任务及按日期查看任务列表。",
    "通知授权、到期提醒、通知中心、通知跳转。",
    "个人资料、通知偏好、外观入口、分类管理、退出登录。",
]:
    add_bullet(doc, text)
add_heading(doc, "1.2 首版明确不做的范围", 2)
for text in [
    "共享任务、成员邀请、评论、@提及、团队空间和协作权限。",
    "循环任务、任务模板、看板、统计报表、桌面/网页端。",
    "深色模式和多语言界面；代码结构会预留扩展点。",
    "任务回收站与复杂的离线双向冲突合并。",
]:
    add_bullet(doc, text)
add_note(doc, "产品原则", "默认以中文简体和用户设备时区运行；任务均为用户私有数据。")

# 2 Requirements
add_page_break(doc)
add_heading(doc, "2. 需求拆解与验收口径")
add_body(doc, "以下内容由线框图转化为可开发、可验收的功能模块。每个模块以可观察的用户结果定义完成标准。")
add_table(doc, ["模块", "功能", "关键验收标准"], [
    ("引导与认证", "欢迎、注册、登录、Google/Apple 登录", "注册字段校验明确；已登录用户重启后直接进入首页；失败可重试。"),
    ("通知授权", "解释提醒价值、请求系统权限、允许跳过", "选择启用时仅请求一次系统权限；无论允许或跳过都能进入首页。"),
    ("首页", "问候、统计、搜索、筛选、任务列表、快速完成", "切换筛选/搜索后结果正确；勾选任务无需进入详情；新增后列表即时刷新。"),
    ("任务详情", "完整字段、子任务、附件、保存、删除", "可编辑并持久化；子任务独立完成；删除前出现确认。"),
    ("创建任务", "标题、备注、日期时间、优先级、分类、提醒", "标题为空不可提交；保存成功后返回任务列表并展示新任务。"),
    ("日历", "月切换、选中日期、当日任务", "有任务的日期有标记；切换月份和日期会正确刷新列表。"),
    ("通知", "到期、完成、延期记录、已读、清除", "通知点击进入对应任务；清除全部后列表和未读数同步更新。"),
    ("设置", "资料、分类、通知偏好、退出", "修改分类会反映到任务表单；退出后清空本地会话并返回欢迎页。")
], [3.0, 5.0, 6.7])
add_heading(doc, "2.1 业务规则默认值", 2)
add_table(doc, ["规则", "默认约定"], [
    ("任务字段", "标题必填；备注、截止时间、提醒、附件与子任务可选。"),
    ("优先级", "低 / 中 / 高，创建时默认“中”。"),
    ("分类", "预置 Work、Personal；用户可新建、重命名和更换颜色，自定义分类不可删除仍被任务引用的记录。"),
    ("提醒时间", "准时、提前 10 分钟、30 分钟、1 小时、1 天；默认关闭。"),
    ("任务状态", "待完成 / 已完成；完成任务保留记录，并显示在“全部”筛选中。"),
    ("删除", "需二次确认；首版直接删除任务、子任务、关联通知与附件引用。"),
    ("附件", "支持图片和 PDF；单文件最大 10 MB，每个任务最多 5 个附件。"),
    ("时间", "以用户设备时区展示；服务端按 UTC 存储，任务记录保存 IANA 时区。")
], [3.6, 11.1])

# 3 Stack and Architecture
add_page_break(doc)
add_heading(doc, "3. 技术方案与架构")
add_body(doc, "采用 Expo 管理的 React Native 应用与 Supabase 后端。该组合能在一套 TypeScript 代码中完成双端界面、认证、数据库、文件与推送集成，同时避免首版维护自建服务器。")
add_table(doc, ["层级", "选型", "使用理由"], [
    ("移动端", "React Native + Expo + TypeScript", "单代码库覆盖 iOS/Android；便于访问通知、文件选择和安全存储。"),
    ("导航", "Expo Router", "按文件组织路由，支持认证守卫、深链及通知跳转。"),
    ("UI", "React Native StyleSheet + 自建设计系统", "线框视觉简洁且定制程度高；避免由通用组件库限制设计。"),
    ("服务端状态", "TanStack Query", "统一请求缓存、失效、加载和错误状态，减少手写同步逻辑。"),
    ("本地 UI 状态", "Zustand", "管理筛选条件、表单草稿和局部交互状态，保持轻量。"),
    ("表单与校验", "React Hook Form + Zod", "将注册、创建任务和设置表单的验证规则集中管理。"),
    ("后端", "Supabase: Auth / PostgreSQL / Storage / Edge Functions", "快速获得认证、行级权限、数据 API、附件存储与受控服务端逻辑。"),
    ("通知", "expo-notifications + Expo Push Service", "统一 iOS/Android Token、权限、提醒接收和点击后的深链处理。"),
    ("日期", "date-fns + date-fns-tz", "处理月历、相对时间、UTC 与用户时区之间的转换。"),
    ("监控", "Sentry", "在测试与发布后收集崩溃、异常与关键通知失败信息。"),
    ("构建发布", "EAS Build / Submit + GitHub Actions", "形成可重复的测试包与商店构建流程。")
], [3.0, 5.2, 6.5])
add_heading(doc, "3.1 逻辑架构", 2)
add_table(doc, ["流向", "职责"], [
    ("Expo App -> Supabase Auth", "注册、邮箱登录、OAuth 回调、会话恢复与退出。"),
    ("Expo App -> Supabase Data API", "受 RLS 保护的任务、分类、子任务、通知读写。"),
    ("Expo App -> Supabase Storage", "以用户目录上传并展示图片/PDF；通过受控访问链接下载。"),
    ("Scheduled Edge Function -> Expo Push", "定期查找应提醒任务，创建通知记录并向有效设备 Token 发推送。"),
    ("点击系统通知 -> 深链 -> 任务详情", "依据 payload 的 taskId 打开详情页并标记通知已读。")
], [5.2, 9.5])
add_note(doc, "通知实施要点", "开发阶段使用 development build 真机验证远程推送；不要以 Expo Go 作为远程推送验收环境。")

# 4 Data security
add_heading(doc, "4. 数据模型、权限与接口边界")
add_heading(doc, "4.1 核心数据表", 2)
add_table(doc, ["表", "关键字段", "用途"], [
    ("profiles", "id, display_name, avatar_url, timezone, preferences", "扩展 auth.users 的用户资料和偏好。"),
    ("categories", "id, user_id, name, color, is_system, sort_order", "系统与自定义任务分类。"),
    ("tasks", "id, user_id, category_id, title, note, status, priority, due_at, reminder_at, timezone", "任务主体；时间均以 UTC 保存。"),
    ("task_subtasks", "id, task_id, title, is_completed, sort_order", "任务的可排序子任务。"),
    ("task_attachments", "id, task_id, storage_path, file_name, mime_type, size_bytes", "任务附件元数据；文件存于 Storage。"),
    ("device_push_tokens", "id, user_id, token, platform, last_seen_at, is_active", "一个用户可有多个设备 Token。"),
    ("notifications", "id, user_id, task_id, type, title, body, is_read, sent_at", "应用内通知中心和推送投递记录。")
], [3.1, 7.0, 4.6])
add_heading(doc, "4.2 Row Level Security 规则", 2)
for text in [
    "profiles：仅本人可查询和更新自己的资料。",
    "categories、tasks、notifications、device_push_tokens：仅允许 user_id = auth.uid() 的行被查询、创建、更新和删除。",
    "task_subtasks、task_attachments：通过关联任务的 user_id 判断归属，禁止直接越权访问。",
    "Storage：路径使用 userId/taskId/filename 结构；仅对象所有者可读写自己的目录。",
    "Edge Function 使用服务端密钥，仅部署在受控环境；移动端绝不保存 service_role key。",
]:
    add_bullet(doc, text)
add_heading(doc, "4.3 接口与动作边界", 2)
add_table(doc, ["动作", "客户端方式", "服务端校验"], [
    ("认证", "Supabase Auth SDK", "邮箱格式、密码长度、OAuth 回调白名单。"),
    ("任务与分类 CRUD", "Supabase 查询 + Zod 预校验", "RLS 归属校验；数据库约束 priority/status 枚举。"),
    ("附件上传", "先校验文件类型/大小，再上传 Storage", "路径所有者、MIME、大小上限；写入元数据。"),
    ("发送提醒", "客户端登记 Token；服务端任务处理", "验证 Token 有效性，幂等写入通知，失败 Token 失活。"),
    ("删除任务", "确认弹窗后提交", "事务清理子任务、通知、附件元数据；异步删除文件。")
], [3.5, 6.4, 4.8])

# 5 UX flows
add_page_break(doc)
add_heading(doc, "5. 关键用户流程与界面实现")
add_table(doc, ["流程", "步骤", "完成条件"], [
    ("新用户进入", "欢迎页 -> 注册/OAuth -> 通知授权 -> 首页", "账户创建、资料初始化、默认分类完成；允许跳过通知授权。"),
    ("创建任务", "首页 + -> 填写任务 -> 创建 -> 首页刷新", "标题为空时阻止提交；成功后定位/展示任务。"),
    ("完成任务", "首页或详情点击复选框 -> 状态更新", "UI 乐观更新；失败时回滚并提示重试。"),
    ("任务提醒", "到达 reminder_at -> 系统推送 -> 点击 -> 详情", "通知 payload 含 taskId；目标任务不存在时回退到通知页。"),
    ("查看日历", "进入日历 -> 选日期或切月份 -> 当日列表刷新", "同一日期结果与首页数据源一致。"),
    ("管理分类", "设置 -> 分类管理 -> 新增/编辑 -> 表单可用", "名称不可为空/重复；被引用分类不允许直接删除。")
], [3.0, 8.4, 3.3])
add_heading(doc, "5.1 页面与路由建议", 2)
add_table(doc, ["路由", "页面", "实现关注点"], [
    ("/(auth)/welcome", "欢迎 / 启动", "加载会话；已登录直接跳转主应用。"),
    ("/(auth)/sign-up, sign-in", "注册 / 登录", "邮箱校验、OAuth、可访问错误提示。"),
    ("/notification-permission", "通知授权", "系统权限仅由用户操作触发；可跳过。"),
    ("/(tabs)/home", "首页任务列表", "搜索防抖、筛选、分页或虚拟列表。"),
    ("/tasks/new, /tasks/[id]", "创建 / 详情", "同一表单模型；详情支持保存、子任务和附件。"),
    ("/(tabs)/calendar", "日历", "按月查询任务概览，按选中日查询列表。"),
    ("/(tabs)/notifications", "通知中心", "未读状态、批量清除、通知深链。"),
    ("/(tabs)/profile", "个人资料 / 设置", "二级路由管理资料、分类、偏好与登出。")
], [4.4, 3.4, 6.9])
add_heading(doc, "5.2 设计系统最小集", 2)
add_body(doc, "先将线框图沉淀为 color、spacing、radius、typography、shadow、icon 和状态色 token。任务卡片、标签、输入框、按钮、空状态、确认弹窗和底部导航应形成可复用组件，避免每个页面单独拼样式。")

# 6 Notification details
add_page_break(doc)
add_heading(doc, "6. 通知与提醒实现方案")
add_body(doc, "提醒既要在用户当前设备上可靠触发，也要在用户换设备、登录多台设备或关闭应用后仍能工作。因此采用“服务端任务生成提醒 + 推送投递 + 应用内通知记录”的方案。")
add_table(doc, ["阶段", "处理逻辑", "失败处理"], [
    ("授权", "用户点击启用后请求系统权限，并为允许的设备获取 Expo Push Token。", "拒绝权限时保留应用内提醒设置入口，不循环弹窗。"),
    ("登记", "登录/启动时 upsert Token、平台、版本和 last_seen_at。", "发送失败或 Token 失效时标记 inactive。"),
    ("调度", "每 5 分钟运行一次服务端任务，找出即将到期且尚未投递的任务。", "用 task_id + reminder_at 唯一键保证幂等，避免重复发送。"),
    ("投递", "写入 notifications，再向该用户所有有效设备发送含 taskId 的推送。", "记录投递回执；临时失败按限定次数重试。"),
    ("点击", "应用读取 payload，路由至 /tasks/[taskId] 并将通知标记已读。", "任务被删除或无权限时跳转通知页并提示。")
], [3.0, 9.2, 3.1])
add_heading(doc, "6.1 通知文案与优先级", 2)
add_table(doc, ["类型", "示例", "行为"], [
    ("到期提醒", "“完成客户方案”将在 30 分钟后到期", "推送 + 应用内通知；点击打开详情。"),
    ("到期", "“完成客户方案”现在到期", "推送 + 应用内通知；高优先级任务可使用更明确文案。"),
    ("完成确认", "“晨间锻炼”已完成", "仅应用内通知，首版不必系统推送。"),
    ("延期", "“取干洗衣物”已延期", "仅应用内通知；创建新版截止提醒。")
], [3.0, 7.7, 4.6])
add_note(doc, "时区注意", "调度以 UTC 比较 reminder_at；界面用任务保存的时区格式化。用户更换设备时区后，既有任务应保持原设定时刻的含义。", RED)

# 7 Delivery plan
add_page_break(doc)
add_heading(doc, "7. 迭代计划与里程碑")
add_body(doc, "以下按 1 个全栈开发者、每个迭代 1 周估算，共 6 周。若由 2 人并行，前端界面与后端/测试可并行推进，但仍应按里程碑验收。")
add_table(doc, ["迭代", "目标与工作内容", "产出 / 验收"], [
    ("第 0 周\n需求冻结", "确认视觉 token、字段规则、提醒默认值、附件策略、账户政策；创建 Supabase 项目与 Apple/Google 配置清单。", "PRD 补充、数据库草案、环境变量清单、可点击流程确认。"),
    ("第 1 周\n基础工程", "Expo/TypeScript、路由、设计系统、Supabase 客户端、认证守卫、CI、错误监控。", "可运行开发包；欢迎、注册、登录、会话恢复可用。"),
    ("第 2 周\n任务核心", "表结构/RLS、分类、首页、创建任务、详情编辑、完成/删除、表单校验。", "个人任务 CRUD 闭环；首页搜索与筛选正确。"),
    ("第 3 周\n增强体验", "子任务、附件上传、日期时间、月历、设置与分类管理、空/加载/错误状态。", "线框图九个主页面可走通；数据刷新一致。"),
    ("第 4 周\n通知", "权限、Token、通知表、服务端调度、Expo Push、深链、通知中心。", "iOS/Android 真机完成准时与提前提醒验收。"),
    ("第 5 周\n测试与发布", "可访问性、边界测试、性能、RLS/安全测试、崩溃监控、商店构建与隐私材料。", "Release Candidate、测试记录、发布构建与回归通过。")
], [2.5, 8.2, 4.6])
add_heading(doc, "7.1 推荐工作分解", 2)
for text in [
    "前端：路由与认证、设计组件、各页面、表单、Query 缓存、推送接收和深链。",
    "后端：数据库迁移、RLS、Storage 策略、Edge Function、通知调度和监控。",
    "测试：单元/组件测试、端到端关键路径、真实设备通知、回归清单。",
    "产品验收：每周对照线框图和本计划的验收口径演示，避免在第 5 周集中发现范围偏差。",
]:
    add_bullet(doc, text)

# 8 Quality
add_page_break(doc)
add_heading(doc, "8. 质量保障、发布与验收")
add_heading(doc, "8.1 测试策略", 2)
add_table(doc, ["层级", "工具", "覆盖重点"], [
    ("单元测试", "Jest", "日期计算、提醒时间、Zod 校验、筛选与排序规则。"),
    ("组件测试", "React Native Testing Library", "表单错误、任务卡片状态、空状态、确认弹窗。"),
    ("端到端", "Maestro", "注册/登录、创建任务、完成任务、日历、退出登录。"),
    ("后端", "SQL 测试脚本 / Supabase 本地环境", "RLS 越权访问、级联删除、通知幂等性。"),
    ("真机验证", "iOS + Android development build", "权限拒绝/允许、远程推送、深链、附件选择与上传。")
], [3.0, 4.4, 8.0])
add_heading(doc, "8.2 发布前验收清单", 2)
for text in [
    "未登录用户不能看到主应用；不同账号之间不能读取、猜测或下载对方数据。",
    "创建、编辑、完成、删除、搜索和筛选在网络较慢/失败时有明确反馈且不会产生错误数据。",
    "日期、月份切换和提醒在常见时区与跨日场景下正确；同一提醒不重复发送。",
    "通知被拒绝后不阻断核心使用；重新开启权限后可恢复 Token 登记。",
    "附件类型、数量和大小限制有前端提示与服务端兜底；删除任务不会留下可访问文件。",
    "iOS/Android 真机的关键流程通过；崩溃与关键接口失败可在监控系统中定位。",
    "应用图标、启动图、隐私政策、数据收集说明、测试账号与商店截图准备完整。",
]:
    add_bullet(doc, text)
add_heading(doc, "8.3 Definition of Done", 2)
add_body(doc, "一个功能只有在实现、代码审查、自动化测试、真机验证、错误状态处理、埋点/日志和产品验收均完成后，才视为完成。没有经过 RLS 验证的数据功能、没有真机验证的推送功能，不进入发布候选版本。")

# 9 Risks
add_page_break(doc)
add_heading(doc, "9. 风险、依赖与开工清单")
add_table(doc, ["风险 / 依赖", "影响", "应对策略"], [
    ("Apple/Google OAuth 配置", "可能影响认证排期", "第 0 周即创建开发者账号、Bundle ID、回调 URL 和测试账号。"),
    ("iOS/Android 推送凭据", "可能导致通知无法真机验收", "第 1 周完成 EAS 项目与凭据配置；第 4 周前先发一条测试推送。"),
    ("系统限制通知准时性", "低电量/系统策略下可能延迟", "以“尽力提醒”设计；服务端推送 + 应用内记录，不承诺秒级准时。"),
    ("附件存储成本", "可能影响长期成本", "首版限制类型、单文件 10 MB、每任务 5 个；后续按实际使用调整。"),
    ("数据隐私与合规", "影响上架与用户信任", "最小化采集、RLS、私有 Storage、明确隐私政策与数据删除流程。"),
    ("需求扩张", "拖慢 MVP", "协作、循环任务、深色模式等进入后续 Backlog，不插入首版关键路径。")
], [4.0, 4.6, 6.8])
add_heading(doc, "9.1 开工前必须准备", 2)
for text in [
    "Supabase 项目、开发/生产环境、数据库迁移仓库与环境变量管理方案。",
    "Apple Developer、Google Cloud OAuth、Expo/EAS 账号及 iOS Bundle ID / Android Package Name。",
    "应用名称、图标、启动图、隐私政策 URL、客服邮箱和应用商店主体信息。",
    "设计源文件或可复用视觉规范：颜色、字体、间距、图标与空状态插图。",
    "测试设备至少各一台 iPhone 与 Android；用于通知、OAuth 和附件流程验收。",
]:
    add_bullet(doc, text)
add_heading(doc, "10. 后续版本 Backlog", 1)
add_table(doc, ["优先级", "候选能力", "进入条件"], [
    ("P1", "循环任务、任务模板、归档/回收站", "MVP 留存与任务完成行为稳定后。"),
    ("P1", "深色模式、多语言、桌面/Web 端", "目标市场及使用设备数据明确后。"),
    ("P2", "共享任务、成员、评论、@提及", "先完成协作角色、邀请和通知策略的独立产品设计。"),
    ("P2", "统计、习惯、AI 任务拆解", "确认用户价值与数据使用边界后。")
], [2.3, 7.0, 6.1])
add_note(doc, "结论", "按本计划推进，可在 6 个一周迭代内交付一个可发布的个人任务管理 MVP；协作功能应在首版数据与体验稳定后作为独立阶段评估。")

doc.save(OUT)
print(OUT.resolve())
